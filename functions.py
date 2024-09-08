from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime
import csv
import generator
import webbrowser
import subprocess
import os
import sys
import shutil
import config

os.chdir('invoice_generator/scripts/')
document = Document('template_inv.docx')
total_amount_of_invoice = 0


def findandreplace(find, replace):
    for paragraph in document.paragraphs:
        if find in paragraph.text:
            paragraph.add_run(replace).bold = True


def add_new_buyer():
    with open('buyer_details.csv', 'r+', newline="") as csvf:
        id_of_buyer = str(len(csvf.readlines()))
        name_of_buyer = input("Podaj nazwę kupującego: ")
        address1_of_buyer = input("Podaj adres ulica plus numer mieszkania: ")
        address2_of_buyer = input("Podaj adres kod pocztowy plus miasto: ")
        nip_of_buyer = input("Podaj nip: ")
        csv_writer = csv.writer(csvf)
        csv_writer.writerow([id_of_buyer, name_of_buyer, address1_of_buyer, address2_of_buyer, nip_of_buyer])


def list_of_buyers():
    with open('buyer_details.csv', newline='') as csvfile:
        csv_reader = csv.DictReader(csvfile)
        for row in csv_reader:
            print(row["id"], row["name"])


def insert_buyer_details():
    while True:
        print("**"*30)
        print("Na kogo chcesz wystawić fakturę:")
        list_of_buyers()
        print("**" * 30)
        buyer_choice = input("\nWybierz opcję: \n  \n 1 - dodaj nowego kontrahenta \n "
                             "2 - wybierz kontrahenta z listy \n \n")
#r - faktura dla rehasport , 3 - edytuj dane wybranego kontrahenta
        if buyer_choice == "1":
            add_new_buyer()
        elif buyer_choice == "2":
            id_of_buyer = input("Podaj id kontrahenta: ")
            with open('buyer_details.csv') as csvfile:
                csv_reader = csv.DictReader(csvfile)
                for row in csv_reader:
                    if row["id"] == id_of_buyer:
                        print("**" * 30)
                        print("\nNAZWA: {} \nADRES: {} {} \nNIP: {}".format(row["name"], row["address1"], row["address2"],
                                                          row["nip"]))
                        print("**" * 30)
                        buyer_name = row["name"]
                        buyer_address1 = row["address1"]
                        buyer_address2 = row["address2"]
                        buyer_nip = row["nip"]
            confirmation = input("\nCzy to ta osoba? y/n: ")
            if confirmation == "y":

                break
            else:
                continue

    document.tables[0].cell(1, 1).text = buyer_name
    document.tables[0].cell(2, 1).text = buyer_address1
    document.tables[0].cell(3, 1).text = buyer_address2
    document.tables[0].cell(4, 1).text = "NIP: {}".format(buyer_nip)
    document.tables[0].cell(1, 1).paragraphs[0].runs[0].font.bold = True

    # with open('buyer_details.csv', newline='') as csvfile:
    #     reader = csv.DictReader(csvfile)
    #     list = list(reader)
    #     for row in list:
    #         print(row)

#         with open("./bwq.csv", 'r') as file:
#             csvreader = csv.reader(file)
#             for row in csvreader:
#                 print(row)
#
# #         if i.get("month") == str(monthnr):
#             y = int(i['nr_inv'])
#             i['nr_inv'] = str(y+1)
#             data = i['nr_inv']
#
# with open('inv_nr.csv', "w", newline='') as csvfile:
#     writer = csv.DictWriter(csvfile, ['month', 'nr_inv'])
#     writer.writeheader()
#     writer.writerows(list)


def set_data_of_invoice():
    while True:
        today = datetime.date.today().strftime("%d/%m/%Y")
        data_choice = input("Czy fakturę wystawić na dzień: {} ? \n\n 1 - tak faktura na dzień dzisiejszy \n "
                            "2 - nie chcę sam podać datę \n".format(today))
        if data_choice == "1":
            findandreplace("Data wystawienia:", today)
            break
        if data_choice == "2":
            data_of_payment = input("Podaj datę w formacie liczbowym dzień/miesiąc/rok: ")
            findandreplace("Data wystawienia:", data_of_payment)
            break
        else:
            print("nierozumiem")

    # monthnr = months.index(month) + 1
    #
    # today = datetime.date.today()
    # year = today.strftime("%Y")
    # if monthnr != 12:
    #     lastday = today.replace(day=1, month=int(monthnr+1), year=int(year)) - datetime.timedelta(days=1)
    # elif monthnr == 12:
    #     lastday = today.replace(day=1, month=1, year=int(year)+1) - datetime.timedelta(days=1)
    # lastdayofmonth = lastday.strftime("%d/%m/%Y")
    # lastdayofmonth2 = lastday.strftime("%m/%Y")
    # term = lastday + datetime.timedelta(days=14)
    # termofpayment = term.strftime("%d/%m/%Y")


def add_item_to_invoice(row):
    name_of_item = input("\nPodaj nazwę usługi/towaru: ")
    number_of_items = input("Podaj ilość wykonanych usług/sprzedanego towaru: ")
    price_for_one_item = input("Podaj kwotę za jedną usługę/towar: ")
    details_of_item = input("Podaj szczegóły do zakupu (np. faktura wystawiona do paragonu z dnia): ")
    amount_of_item = int(number_of_items) * int(price_for_one_item)
    global total_amount_of_invoice
    total_amount_of_invoice += amount_of_item

    new_row = document.tables[1].add_row().cells
    new_row[0].text = str(row)
    new_row[1].text = name_of_item
    new_row[2].text = "szt."
    new_row[3].text = number_of_items
    new_row[4].text = price_for_one_item
    new_row[5].text = "zw"
    new_row[6].text = str(amount_of_item)
    new_row[7].text = "0,00"
    new_row[8].text = str(amount_of_item)
    new_row[9].text = details_of_item

    document.tables[1].cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.tables[1].cell(row, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.tables[1].cell(row, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.tables[1].cell(row, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.tables[1].cell(row, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.tables[1].cell(row, 6).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.tables[1].cell(row, 7).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.tables[1].cell(row, 8).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.tables[1].cell(row, 9).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER


def insert_seller_details():
    document.tables[0].cell(1, 0).text = config.name_1
    document.tables[0].cell(2, 0).text = config.adress_1
    document.tables[0].cell(3, 0).text = config.additional_adress_1
    document.tables[0].cell(4, 0).text = "NIP: " + config.nip_1
    document.tables[0].cell(5, 0).text = "E-mail: " + config.mail_1
    document.tables[0].cell(6, 0).text = "Tel.: " + config.tel_1
    document.tables[0].cell(1, 0).paragraphs[0].runs[0].font.bold = True


def insert_total_amount():
    global total_amount_of_invoice
    total_amount_of_invoice_str = str(total_amount_of_invoice)
    document.tables[2].cell(1, 1).text = total_amount_of_invoice_str
    document.tables[2].cell(1, 3).text = total_amount_of_invoice_str
    document.tables[2].cell(2, 1).text = total_amount_of_invoice_str
    document.tables[2].cell(2, 3).text = total_amount_of_invoice_str
    document.tables[2].cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.tables[2].cell(1, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.tables[2].cell(2, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.tables[2].cell(2, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER


def payment_information():
    global total_amount_of_invoice
    total_amount_of_invoice_str = str(total_amount_of_invoice)
    document.tables[3].cell(0, 3).text = total_amount_of_invoice_str + " PLN"
    # document.tables[3].cell(2, 3).text = total_amount_of_invoice_str+" PLN"
    # document.tables[3].cell(3, 3).text = word_value
    # document.tables[3].cell(1 ,1).text = termofpayment
    # document.tables[3].cell(2, 1).text = config.bank_name
    # document.tables[3].cell(3, 1).text = config.bank_account


def save_as_docx():
    name_of_file = "faktura nr ..."
    document.save(name_of_file + ".docx")


def convert_to_pdf():
    name_of_file = "faktura nr ..."
    output = subprocess.check_output(['libreoffice', '--convert-to', 'pdf', name_of_file + ".docx"])
    webbrowser.open(name_of_file + ".pdf", new=1)


def move_file():
    dest_fpath = config.path
    shutil.move("faktura nr ....pdf", dest_fpath+"faktura nr....pdf")


#####################################################################################

####################################################################################
# wprowadź kwotę faktury
# dane kupującego
#     lista wyboru albo wprowadź nowego
# nazwa usługi
# ilość
# cena za jedną sztukę


# os.chdir(os.path.dirname(sys.argv[0]))
#
# months =["styczeń", "luty", "marzec", "kwiecień", "maj", "czerwiec", "lipiec", "sierpień", "wrzesień", "październik", "listopad", "grudzień"]
#
# print("Program do faktur")
#
# while True:
#     value = input("Podaj kwotę pieniedzy na jaką ma być wystawiona faktura: ")
#     value = value.replace(",",".")
#     try:
#         float(value)
#     except ValueError:
#         continue
#     break
#
# while True:
#     month = input("Podaj miesiąc w którym wykonana była usługa:")
#     month = month.lower()
#     if month in months:
#         break
#     else:
#         print("Nie ma takiego miesiąca spróbuj jeszcze raz:")
#
# ########################################################################################
#
#
#
# name = "usługi medyczne wykonane w miesiącu - "+ month + " "+year+""
# netto = value


# with open('inv_nr.csv', newline='') as csvfile:
#     reader = csv.DictReader(csvfile)
#     list = list(reader)
#     for i in list:
#         if i.get("month") == str(monthnr):
#             y = int(i['nr_inv'])
#             i['nr_inv'] = str(y+1)
#             data = i['nr_inv']
#
# with open('inv_nr.csv', "w", newline='') as csvfile:
#     writer = csv.DictWriter(csvfile, ['month', 'nr_inv'])
#     writer.writeheader()
#     writer.writerows(list)
#
# word_value = str(generator.generator(value))
# #############################################################################
#
# def findandreplace(find, replace):
#     for paragraph in document.paragraphs:
#         if find in paragraph.text:
#             paragraph.add_run(replace).bold = True


# document = Document('template_inv.docx')
# print(data)
# findandreplace("Faktura nr FV", ""+data+"/" + lastdayofmonth2)
# findandreplace("Data wystawienia:", lastdayofmonth)
# findandreplace("Miejsce wystawienia", config.place)
#
# ###################### TABLE1 #####################
#
#
#
#
#
# ###################### TABLE2 #####################
# document.tables[1].cell(1, 1).text = name
#
# document.tables[1].cell(1, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
# document.tables[1].cell(1, 7).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
# document.tables[1].cell(1, 9).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
#
# ###################### TABLE3 #####################
#
#






# while True:
#     approve = input("Czy faktura jest ok? y/n: ")
#     approve = approve.lower()
#     if approve == "y":
#         os.makedirs(os.path.dirname(dest_fpath+month+"/"), exist_ok=True)
#         shutil.move(nameoffile+".docx", dest_fpath+month+"/"+nameoffile+".docx")
#         shutil.move(nameoffile+".pdf", dest_fpath+month+"/"+nameoffile+".pdf")
#         break
#     elif approve == "n":
#         while True:
#             x = input("Wybierz opcję: \ndel - usuń plik \nedit - zostaw plik .docx: ")
#             if x == "del":
#                 os.remove(nameoffile+".docx")
#                 os.remove(nameoffile+".pdf")
#                 print("Usunięto wszystkie pliki")
#                 break
#             if x == "edit":
#                 os.remove(nameoffile+".pdf")
#                 print("Popraw błędy w pliku .docx")
#                 break
#             else:
#                 continue
#         break
#     else:
#         continue
